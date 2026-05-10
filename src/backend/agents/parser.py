"""Parser Agent — 多格式教材解析与结构化。

Supports: PDF (PyMuPDF), Markdown, TXT.
Strategy: page-by-page parsing for memory efficiency on large files (826MB total).
"""
from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF

from models.schemas import ChapterOut, TextbookOut

# ── Chapter detection ──────────────────────────────────

# Matches: "第一章", "第 1 章", "第12章", "第十二章"
CHAPTER_RE = re.compile(
    r"第\s*([一二三四五六七八九十百零\d]+)\s*章\s*[\s\S]*",
)

# Matches: "Chapter 1", "Chapter 12"
CHAPTER_EN_RE = re.compile(
    r"Chapter\s+(\d+)\s*[\s\S]*",
    re.IGNORECASE,
)

# Chinese numerals → int mapping (up to 50, covers most textbooks)
_CN_NUM = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
    "十一": 11, "十二": 12, "十三": 13, "十四": 14, "十五": 15,
    "十六": 16, "十七": 17, "十八": 18, "十九": 19, "二十": 20,
    "二十一": 21, "二十二": 22, "二十三": 23, "二十四": 24,
    "二十五": 25, "二十六": 26, "二十七": 27, "二十八": 28,
    "二十九": 29, "三十": 30,
    "三十一": 31, "三十二": 32, "三十三": 33, "三十四": 34, "三十五": 35,
    "三十六": 36, "三十七": 37, "三十八": 38, "三十九": 39, "四十": 40,
    "四十一": 41, "四十二": 42, "四十三": 43, "四十四": 44, "四十五": 45,
    "四十六": 46, "四十七": 47, "四十八": 48, "四十九": 49, "五十": 50,
    "五十一": 51, "五十二": 52, "五十三": 53, "五十四": 54, "五十五": 55,
    "五十六": 56, "五十七": 57, "五十八": 58, "五十九": 59, "六十": 60,
}

# Lines that look like headers/footers (page numbers, repeated boilerplate)
_FOOTER_PATTERNS = [
    re.compile(r"^\s*\d+\s*$"),
    re.compile(r"^第\s*\d+\s*页"),
    re.compile(r"^版权所有"),
    re.compile(r"^\s*$"),
]


def _cn_to_int(cn: str) -> int:
    """Convert Chinese numeral to int. Supports 一 through 六十."""
    if cn in _CN_NUM:
        return _CN_NUM[cn]
    if cn.isdigit():
        return int(cn)
    return 0


def _looks_like_header_footer(text: str) -> bool:
    """Heuristic: short line that matches known footer patterns."""
    text = text.strip()
    if len(text) < 30 and any(p.search(text) for p in _FOOTER_PATTERNS):
        return True
    return False


def _extract_chapter_number(text: str) -> Optional[int]:
    """Try to extract a chapter number from a heading line. Returns None on failure."""
    m = CHAPTER_RE.match(text.strip())
    if m:
        return _cn_to_int(m.group(1))
    m = CHAPTER_EN_RE.match(text.strip())
    if m:
        return _cn_to_int(m.group(1))
    return None


# ── PDF Parsing ────────────────────────────────────────

def _extract_text_with_fonts(page: fitz.Page) -> list[dict]:
    """Extract text blocks from a single page, each with font size metadata."""
    blocks = page.get_text("dict")["blocks"]
    lines: list[dict] = []
    for block in blocks:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            text = "".join(s["text"] for s in spans)
            sizes = [s["size"] for s in spans if "size" in s]
            font_name = spans[0].get("font", "")
            avg_size = sum(sizes) / len(sizes) if sizes else 11.0
            is_bold = "bold" in font_name.lower() or "Bold" in font_name
            y0 = line["bbox"][1]
            lines.append({
                "text": text,
                "size": round(avg_size, 1),
                "is_bold": is_bold,
                "y": y0,
            })
    return lines


def _find_body_font_size(all_lines: list[dict]) -> float:
    """Estimate the body-text font size by taking the median font size."""
    sizes = sorted(ln["size"] for ln in all_lines if len(ln["text"].strip()) > 10)
    if not sizes:
        return 11.0
    return sizes[len(sizes) // 2]


def _parse_pdf(file_path: str) -> TextbookOut:
    """Parse a PDF textbook page by page.

    Strategy:
      1. Scan all pages, collecting lines with font metadata.
      2. Identify chapter boundaries: regex + font size heuristic.
      3. Group pages into chapters.
      4. Filter header/footer lines.
    """
    doc = fitz.open(file_path)
    filename = Path(file_path).name
    raw_title = Path(file_path).stem
    # Strip numeric prefix like "01_" for display
    title = re.sub(r"^\d{2,3}[_\-]", "", raw_title)

    page_lines: list[list[dict]] = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        lines = _extract_text_with_fonts(page)
        page_lines.append(lines)
    doc.close()

    all_lines_flat = [ln for pg in page_lines for ln in pg]
    body_size = _find_body_font_size(all_lines_flat)
    size_threshold = body_size * 1.15

    chapters: list[dict] = []
    for pg_idx, lines in enumerate(page_lines):
        for ln in lines:
            text = ln["text"].strip()
            ch_num = _extract_chapter_number(text)
            if ch_num is not None:
                if ln["size"] >= size_threshold or ln["is_bold"] or len(text) < 40:
                    if chapters and chapters[-1]["start_page"] == pg_idx:
                        chapters[-1]["title"] = text
                    else:
                        chapters.append({"title": text, "start_page": pg_idx})
                    break

    total_pages = len(page_lines)
    chapter_objs: list[ChapterOut] = []

    if not chapters:
        chapter_objs.append(ChapterOut(
            chapter_id="ch_01", title=title,
            page_start=1, page_end=total_pages,
            content="", char_count=0,
        ))
    else:
        for i, ch in enumerate(chapters):
            ch_start = ch["start_page"]
            ch_end = chapters[i + 1]["start_page"] - 1 if i + 1 < len(chapters) else total_pages - 1
            chapter_objs.append(ChapterOut(
                chapter_id=f"ch_{i + 1:02d}",
                title=ch["title"],
                page_start=ch_start + 1,
                page_end=ch_end + 1,
                content="",
                char_count=0,
            ))

    # Extract content per chapter
    for ch_obj in chapter_objs:
        start_pg = ch_obj.page_start - 1
        end_pg = ch_obj.page_end - 1
        content_parts: list[str] = []
        for pg_idx in range(start_pg, end_pg + 1):
            lines = page_lines[pg_idx]
            body_lines = [
                ln["text"] for ln in lines
                if not _looks_like_header_footer(ln["text"])
            ]
            content_parts.extend(body_lines)
        ch_obj.content = "\n".join(content_parts)
        ch_obj.char_count = len(ch_obj.content.replace("\n", "").replace(" ", ""))

    total_chars = sum(c.char_count for c in chapter_objs)

    return TextbookOut(
        textbook_id=f"book_{hashlib.md5(filename.encode()).hexdigest()[:8]}",
        filename=filename,
        title=title,
        total_pages=total_pages,
        total_chars=total_chars,
        chapters=chapter_objs,
    )


# ── Markdown / TXT ─────────────────────────────────────

_MD_HEADING_RE = re.compile(r"^#{1,3}\s+(.+)$")


def _parse_markdown(file_path: str) -> TextbookOut:
    filename = Path(file_path).name
    title = re.sub(r"^\d{2,3}[_\-]", "", Path(file_path).stem)
    content = Path(file_path).read_text(encoding="utf-8")
    lines = content.split("\n")
    chapters: list[ChapterOut] = []
    current_title = "前言"
    current_lines: list[str] = []
    ch_idx = 0
    for line in lines:
        m = _MD_HEADING_RE.match(line)
        if m:
            if current_lines:
                ch_idx += 1
                body = "\n".join(current_lines)
                chapters.append(ChapterOut(
                    chapter_id=f"ch_{ch_idx:02d}", title=current_title,
                    page_start=1, page_end=1, content=body,
                    char_count=len(body.replace("\n", "").replace(" ", "")),
                ))
            current_title = m.group(1).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_lines:
        ch_idx += 1
        body = "\n".join(current_lines)
        chapters.append(ChapterOut(
            chapter_id=f"ch_{ch_idx:02d}", title=current_title,
            page_start=1, page_end=1, content=body,
            char_count=len(body.replace("\n", "").replace(" ", "")),
        ))
    if not chapters:
        chapters.append(ChapterOut(
            chapter_id="ch_01", title=title, page_start=1, page_end=1,
            content=content,
            char_count=len(content.replace("\n", "").replace(" ", "")),
        ))
    total_chars = sum(c.char_count for c in chapters)
    return TextbookOut(
        textbook_id=f"book_{hashlib.md5(filename.encode()).hexdigest()[:8]}",
        filename=filename, title=title,
        total_pages=1, total_chars=total_chars, chapters=chapters,
    )


def _parse_txt(file_path: str) -> TextbookOut:
    filename = Path(file_path).name
    title = re.sub(r"^\d{2,3}[_\-]", "", Path(file_path).stem)
    content = Path(file_path).read_text(encoding="utf-8")
    clean = content.replace("\n", "").replace(" ", "")
    return TextbookOut(
        textbook_id=f"book_{hashlib.md5(filename.encode()).hexdigest()[:8]}",
        filename=filename, title=title,
        total_pages=1, total_chars=len(clean),
        chapters=[ChapterOut(
            chapter_id="ch_01", title=title, page_start=1, page_end=1,
            content=content, char_count=len(clean),
        )],
    )


# ── Word (.docx) ────────────────────────────────────────

def _parse_docx(file_path: str) -> TextbookOut:
    """Parse a Word document. Uses heading styles to detect chapters."""
    from docx import Document

    doc = Document(file_path)
    filename = Path(file_path).name
    title = re.sub(r"^\d{2,3}[_\-]", "", Path(file_path).stem)

    chapters: list[ChapterOut] = []
    current_title = title
    current_paras: list[str] = []
    ch_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Detect heading by style name (Heading 1, Heading 2, etc.)
        style_name = para.style.name if para.style else ""
        is_heading = (
            style_name.startswith("Heading") or
            style_name.startswith("heading") or
            "标题" in style_name
        )
        if is_heading:
            if current_paras:
                ch_idx += 1
                body = "\n".join(current_paras)
                chapters.append(ChapterOut(
                    chapter_id=f"ch_{ch_idx:02d}", title=current_title,
                    page_start=1, page_end=1, content=body,
                    char_count=len(body.replace("\n", "").replace(" ", "")),
                ))
            current_title = text
            current_paras = []
        else:
            current_paras.append(text)

    if current_paras:
        ch_idx += 1
        body = "\n".join(current_paras)
        chapters.append(ChapterOut(
            chapter_id=f"ch_{ch_idx:02d}", title=current_title,
            page_start=1, page_end=1, content=body,
            char_count=len(body.replace("\n", "").replace(" ", "")),
        ))

    if not chapters:
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        clean = all_text.replace("\n", "").replace(" ", "")
        chapters.append(ChapterOut(
            chapter_id="ch_01", title=title, page_start=1, page_end=1,
            content=all_text, char_count=len(clean),
        ))

    total_chars = sum(c.char_count for c in chapters)
    return TextbookOut(
        textbook_id=f"book_{hashlib.md5(filename.encode()).hexdigest()[:8]}",
        filename=filename, title=title,
        total_pages=1, total_chars=total_chars, chapters=chapters,
    )


# ── Excel (.xlsx) ───────────────────────────────────────

def _parse_xlsx(file_path: str) -> TextbookOut:
    """Parse an Excel file. Each sheet becomes a chapter."""
    import pandas as pd

    filename = Path(file_path).name
    title = Path(file_path).stem
    sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)

    chapters: list[ChapterOut] = []
    for sheet_name, df in sheets.items():
        if df.empty:
            continue
        # Build content: column headers + each row as text
        cols = list(df.columns.astype(str))
        rows_text: list[str] = [", ".join(cols)]
        for _, row in df.iterrows():
            row_vals = [str(v) for v in row.values if str(v) not in ("nan", "None", "")]
            if row_vals:
                rows_text.append(" | ".join(row_vals))

        body = "\n".join(rows_text)
        clean = body.replace("\n", "").replace(" ", "")
        chapters.append(ChapterOut(
            chapter_id=f"ch_{len(chapters) + 1:02d}",
            title=f"{title} - {sheet_name}",
            page_start=1, page_end=1,
            content=body,
            char_count=len(clean),
        ))

    if not chapters:
        all_dfs = pd.read_excel(file_path, sheet_name=None, dtype=str)
        all_rows = []
        for _, df in all_dfs.items():
            all_rows.extend(df.astype(str).to_csv(index=False).split("\n"))
        content = "\n".join(all_rows)
        clean = content.replace("\n", "").replace(" ", "")
        chapters.append(ChapterOut(
            chapter_id="ch_01", title=title, page_start=1, page_end=1,
            content=content, char_count=len(clean),
        ))

    total_chars = sum(c.char_count for c in chapters)
    return TextbookOut(
        textbook_id=f"book_{hashlib.md5(filename.encode()).hexdigest()[:8]}",
        filename=filename, title=title,
        total_pages=1, total_chars=total_chars, chapters=chapters,
    )


# ── Public API ─────────────────────────────────────────

SUPPORTED_FORMATS = {".pdf", ".md", ".txt", ".markdown", ".docx", ".xlsx"}

_FORMAT_PARSERS = {
    ".pdf": _parse_pdf,
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".txt": _parse_txt,
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
}


async def parse_textbook(file_path: str, filename: Optional[str] = None) -> TextbookOut:
    if filename is None:
        filename = Path(file_path).name
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format '{ext}'. Supported: {SUPPORTED_FORMATS}")
    parser = _FORMAT_PARSERS[ext]
    return parser(file_path)
